import streamlit as st
import os
import logging
from dotenv import load_dotenv
import pandas as pd
from docx import Document as DocxDocument  # pip install python-docx
from langchain.schema import Document as LangchainDocument
from langchain.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import ConversationalRetrievalChain
import requests

# Enable debug logging
logging.basicConfig(level=logging.DEBUG)

# --- Load API keys from Streamlit secrets ---
GROQ_API_KEY = st.secrets.get("GROQ_API_KEY", "").strip()
SERPER_API_KEY = st.secrets.get("SERPER_API_KEY", "").strip()

# --- Check if both keys are found ---
if not GROQ_API_KEY:
    st.error("GROQ_API_KEY not found in Streamlit secrets. Please set it in your app's secrets configuration.")
    st.stop()

if not SERPER_API_KEY:
    st.error("SERPER_API_KEY not found in Streamlit secrets. Please set it in your app's secrets configuration.")
    st.stop()

# --- Set as environment variables if needed ---
os.environ["GROQ_API_KEY"] = GROQ_API_KEY
os.environ["SERPER_API_KEY"] = SERPER_API_KEY

# Page config (must be set before any other Streamlit commands)
st.set_page_config(
    page_title="Denmark Competitors Q&A",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Validate keys presence
if not GROQ_API_KEY:
    st.error("GROQ_API_KEY not found in environment. Please set it in API_KEY.env without extra spaces or comments.")
    st.stop()
if not SERPER_API_KEY:
    st.error("SERPER_API_KEY not found in environment. Please set it in API_KEY.env without extra spaces or comments.")
    st.stop()

# Initialize chat history
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []

# --- Sidebar info ---
with st.sidebar:
    st.title("About")
    st.markdown(
        "Identify potential competitors among high-growth firms in Denmark.\nPowered by Groq AI and LangChain."
    )
    st.markdown("---")
    st.sidebar.markdown(
     """
    **Model Selection**  
    - **llama-3.3-70b-versatile**: Deeper, more nuanced answers; handles longer documents better.   
    - **qwen-qwq-32b**: Faster, lighter, and cheaper; ideal for quick or smaller tasks.  

    **Note:** Llama can be 2×–3× more expensive than Qwen for similar tasks.
    """
)
    model_choice = st.selectbox(
        "Select LLM model:",
        options=["llama-3.3-70b-versatile", "qwen-qwq-32b"],
        index=0,
        key="model_choice"
    )
    st.markdown("---")
    st.markdown("### Tips")
    st.markdown(
        "- Ask specific competitor questions\n"
        "- Include context for better answers\n"
        "- Toggle web search for fresh data"
    )
    st.markdown("---")
    use_web = st.checkbox(
        "Enable Web Search",
        value=True,
        key="enable_web_search"
    )

# Main header
st.title("Denmark Competitors Q&A Assistant")
st.markdown("Ask about high-growth Danish firms and get AI-powered competitor insights.")

# Load and build chain, parameterized on model name
def load_documents():
    df = pd.read_excel('V5_denmark_companies_with_merged_topics.xlsx')
    df = df.drop(columns=[c for c in df.columns if 'Unnamed' in c or df[c].isnull().all()])
    df.fillna('', inplace=True)
    df['combined_text'] = (
        df['Company name Latin alphabet'] + ' - ' +
        df['City Latin Alphabet'] + ' ' +
        df['Region in country'] + ' ' +
        df['Standardized legal form'] + ' ' +
        df['Final Description'] + ' ' +
        df['Topic - Umbrella (Merged)']
    )
    return [
        LangchainDocument(
            page_content=row['combined_text'],
            metadata={
                'company': row['Company name Latin alphabet'],
                'founded_year': row['Founded Year'],
                'sector': row['NACE Rev. 2 main section'],
                'n_emp': row['Number of employees 2023'],
                'growth': row['Growth 2023']
            }
        ) for _, row in df.iterrows()
    ]

@st.cache_resource
def build_chain(model_name: str):
    docs = load_documents()
    embeddings = HuggingFaceEmbeddings(
        model_name='sentence-transformers/all-MiniLM-L6-v2',
        model_kwargs={'device':'cpu'}
    )
    vectorstore = FAISS.from_documents(docs, embeddings)

    llm = ChatGroq(
        model=model_name,
        api_key=GROQ_API_KEY,
        timeout=30,
        max_retries=3
    )

    prompt = PromptTemplate.from_template("""
You are an expert business analyst.
Given context on Danish high-growth firms, return the top 3 competitors for the user's query, each with bullet-point reasons.

Context:
{context}

Question:
{question}
""")

    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(),
        combine_docs_chain_kwargs={'prompt': prompt},
        verbose=True
    )

# Build chain for selected model (rebuilds when model_choice changes)
chain = build_chain(model_choice)

# Display chat history
def render_history():
    for msg in st.session_state.chat_history:
        with st.chat_message(msg['role']):
            st.markdown(msg['content'])
render_history()

# Web search helper
def serper_search(q):
    try:
        r = requests.post(
            'https://google.serper.dev/search',
            headers={'X-API-KEY': SERPER_API_KEY},
            json={'q': q, 'num': 3}
        )
        r.raise_for_status()
        return [
            f"- {it['title']}: {it.get('snippet','')} ({it['link']})"
            for it in r.json().get('organic', [])
        ]
    except Exception as e:
        return [f"Web search failed: {e}"]

# Answer generation
if prompt := st.chat_input("Ask your question about competitors..."):
    st.session_state.chat_history.append({'role':'user','content':prompt})
    with st.chat_message('user'):
        st.markdown(prompt)
    with st.chat_message('assistant'):
        with st.spinner('Thinking...'):
            context = prompt
            if use_web:
                web_results = serper_search(prompt)
                context += '\n\nWeb Search Results:\n' + '\n'.join(web_results)
            try:
                out = chain.invoke({'question': context, 'chat_history': []})
            except Exception as e:
                st.error(f"Groq API error: {e}")
                st.stop()
            ans = out.get('answer','')
            st.markdown(ans)
            st.session_state.chat_history.append({'role':'assistant','content':ans})

            # Download report
            if ans:
                fp = 'report.docx'
                doc = DocxDocument()
                doc.add_heading('Competitor Analysis Report', level=1)
                for line in ans.splitlines():
                    doc.add_paragraph(line)
                doc.save(fp)
                with open(fp, 'rb') as f:
                    st.download_button(
                        'Download Report',
                        f,
                        'competitor_analysis_report.docx',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )

# Clear history button
with st.sidebar:
    if st.button('Clear Chat History', key='clear_history'):
        st.session_state.pop('chat_history', None)

# Footer
st.markdown('---')
st.markdown(
    """
<div style='text-align:center'>
  <p>Built using Streamlit and LangChain</p>
  <p>Powered by Groq AI</p>
</div>
""",
    unsafe_allow_html=True
)